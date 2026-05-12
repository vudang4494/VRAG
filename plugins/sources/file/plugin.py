"""File source plugin — PDF, DOCX, TXT, CSV, Excel, Markdown."""
import asyncio
import hashlib
import io
from datetime import datetime
from pathlib import Path
from typing import Any

from loguru import logger

from plugins.base import (
    BaseSourcePlugin,
    ParsedDocument,
    PluginCapability,
    PluginConfig,
    SourceCredentials,
)


class FileSourcePlugin(BaseSourcePlugin):
    """Ingest local/remote files via upload or URL."""

    name = "file"
    version = "1.0.0"
    capabilities = [
        PluginCapability.INGEST_FILE,
        PluginCapability.INGEST_URL,
        PluginCapability.INGEST_STREAM,
    ]
    supported_types = ["pdf", "docx", "doc", "txt", "md", "markdown", "csv", "xlsx", "xls"]

    def _parse_pdf(self, content: bytes, filename: str) -> str:
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert_bytes(content, source_type=Path(filename).suffix.lstrip("."))
            return result.document.export_to_markdown()
        except ImportError:
            pass

        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
        except ImportError:
            pass

        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            texts = [doc[i].get_text() for i in range(len(doc))]
            return "\n\n".join(texts)
        except ImportError:
            pass

        logger.warning(f"PDF parsing failed for {filename}, no PDF library available")
        return ""

    def _parse_docx(self, content: bytes, filename: str) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            pass

        try:
            import mammoth
            result = mammoth.extract_raw_text(io.BytesIO(content))
            return result.value
        except ImportError:
            pass

        return ""

    def _parse_csv(self, content: bytes, filename: str) -> str:
        try:
            import csv
            import io as _io
            reader = csv.reader(_io.TextIOWrapper(_io.BytesIO(content)))
            rows = [" | ".join(row) for row in reader if any(row)]
            return "\n".join(rows)
        except Exception:
            return content.decode("utf-8", errors="replace")

    def _parse_excel(self, content: bytes, filename: str) -> str:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"## Sheet: {sheet}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except ImportError:
            pass

        try:
            import pandas as pd
            df = pd.read_excel(content)
            return df.to_string()
        except ImportError:
            pass

        return content.decode("utf-8", errors="replace")

    def _parse_markdown(self, content: bytes, filename: str) -> str:
        return content.decode("utf-8", errors="replace")

    def _parse_txt(self, content: bytes, filename: str) -> str:
        return content.decode("utf-8", errors="replace")

    def _parse_by_extension(self, content: bytes, filename: str) -> str:
        ext = Path(filename).suffix.lower().lstrip(".")
        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "doc": self._parse_docx,
            "csv": self._parse_csv,
            "xlsx": self._parse_excel,
            "xls": self._parse_excel,
            "md": self._parse_markdown,
            "markdown": self._parse_markdown,
            "txt": self._parse_txt,
        }
        parser = parser_map.get(ext, self._parse_txt)
        text = parser(content, filename)
        return self._normalize_text(text)

    async def fetch(self, url_or_path: str, **kwargs: Any) -> ParsedDocument:
        content = kwargs.get("content", b"")
        filename = kwargs.get("filename", url_or_path or "document.txt")
        metadata = kwargs.get("metadata", {})

        text = self._parse_by_extension(content, filename)

        doc_hash = hashlib.md5(content).hexdigest()[:16]

        return ParsedDocument(
            title=metadata.get("title", Path(filename).stem),
            content=text,
            raw_content=content if len(content) < 10 * 1024 * 1024 else None,
            url=url_or_path if url_or_path.startswith("http") else None,
            author=metadata.get("author"),
            created_date=metadata.get("created_date"),
            file_type=Path(filename).suffix.lower().lstrip("."),
            file_size_bytes=len(content),
            metadata={
                **metadata,
                "doc_hash": doc_hash,
                "ingested_via": "file_plugin",
            },
        )

    async def ingest(self, url_or_path: str, **kwargs: Any) -> Any:
        doc = await self.fetch(url_or_path, **kwargs)
        from src.storage.document_store import DocumentStore
        store = DocumentStore()
        return await store.ingest_document(
            tenant_id=self.config.tenant_id,
            source_id=self.config.source_id,
            doc=doc,
            chunk_size=self.config.get("chunk_size", 512),
            chunk_overlap=self.config.get("chunk_overlap", 64),
        )
