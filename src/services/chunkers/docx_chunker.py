"""DOCX chunker — heading-hierarchy aware.

Strategy: parse document → emit sections per Heading 1-3 → emit paragraphs.
Preserves heading_path metadata.
"""
from __future__ import annotations

from typing import Any

from loguru import logger

from src.services.chunkers.base import BaseChunker, ChunkUnit


class DocxChunker(BaseChunker):
    name = "docx"

    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    async def chunk(self, content: bytes | str, filename: str = "") -> list[ChunkUnit]:
        if isinstance(content, str):
            content = content.encode("utf-8")

        sections = self._parse_with_headings(content)
        if not sections:
            return []

        units: list[ChunkUnit] = []
        idx = 0

        for sec_idx, (heading_path, paragraphs) in enumerate(sections):
            section_text = "\n\n".join(paragraphs)
            section_parent = None
            if "section" in self.emit_levels:
                units.append(ChunkUnit(
                    text=section_text,
                    chunk_index=idx,
                    chunk_level="section",
                    parent_index=None,
                    metadata={
                        "heading_path": heading_path,
                        "section_index": sec_idx,
                        "filename": filename,
                        "format": "docx",
                    },
                ))
                section_parent = idx
                idx += 1

            # Pack paragraphs để không vượt paragraph_max_chars
            packed = self.pack_units(paragraphs, self.paragraph_max_chars, joiner="\n\n")
            for p_idx, para in enumerate(packed):
                if "paragraph" not in self.emit_levels:
                    continue
                units.append(ChunkUnit(
                    text=para,
                    chunk_index=idx,
                    chunk_level="paragraph",
                    parent_index=section_parent,
                    metadata={
                        "heading_path": heading_path,
                        "section_index": sec_idx,
                        "paragraph_index": p_idx,
                        "filename": filename,
                        "format": "docx",
                    },
                ))
                idx += 1

        return units

    def _parse_with_headings(self, content: bytes) -> list[tuple[list[str], list[str]]]:
        """Return list of (heading_path, paragraph_texts) per section."""
        try:
            from io import BytesIO
            from docx import Document
        except ImportError:
            return self._parse_mammoth_fallback(content)

        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            logger.debug(f"python-docx parse failed, fallback: {e}")
            return self._parse_mammoth_fallback(content)

        sections: list[tuple[list[str], list[str]]] = []
        heading_stack: list[str] = []
        current_paras: list[str] = []

        def _flush():
            if current_paras:
                sections.append((heading_stack.copy(), current_paras.copy()))
                current_paras.clear()

        for para in doc.paragraphs:
            style = (para.style.name if para.style else "") or ""
            txt = (para.text or "").strip()
            if not txt:
                continue
            if style.startswith("Heading"):
                _flush()
                try:
                    level = int(style.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
                # Trim stack to current level, then push
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(txt)
            else:
                current_paras.append(txt)
        _flush()

        # Cover tables
        try:
            for t_idx, table in enumerate(doc.tables):
                rows_md = []
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    rows_md.append("| " + " | ".join(cells) + " |")
                if rows_md:
                    table_md = "\n".join(rows_md)
                    sections.append((heading_stack.copy() + [f"Table {t_idx + 1}"], [table_md]))
        except Exception:
            pass

        return sections or self._parse_mammoth_fallback(content)

    def _parse_mammoth_fallback(self, content: bytes) -> list[tuple[list[str], list[str]]]:
        try:
            import mammoth
            from io import BytesIO

            result = mammoth.extract_raw_text(BytesIO(content))
            txt = result.value or ""
            paragraphs = [p.strip() for p in txt.split("\n\n") if p.strip()]
            return [([], paragraphs)] if paragraphs else []
        except Exception as e:
            logger.warning(f"mammoth fallback also failed: {e}")
            return []
